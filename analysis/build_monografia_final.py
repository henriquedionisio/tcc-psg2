#!/usr/bin/env python3
"""Gera monografia-final.docx a partir de monografia-final.md, preenchendo as
páginas pré-textuais do template EACH (monografia-modelo.docx) e embutindo as
figuras principais no corpo."""

from __future__ import annotations

import sys
from pathlib import Path

ROOT = Path(__file__).resolve().parent.parent
sys.path.insert(0, str(ROOT))

import re

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt
from docx.text.paragraph import Paragraph

TEMPLATE_PATH = ROOT / "docs" / "monografia" / "monografia-modelo.docx"
MD_PATH = ROOT / "docs" / "monografia" / "monografia-final.md"
OUTPUT_PATH = ROOT / "docs" / "monografia" / "monografia-final.docx"

FIGURE_COUNTER = 0

TITULO = (
    "Avaliação de Variabilidade em Large Language Models "
    "por meio de Digital Twins Conversacionais"
)
AUTOR = "HENRIQUE COSTA DIONÍSIO"

RESUMO_REF = (
    "DIONÍSIO, Henrique Costa. Avaliação de Variabilidade em Large Language "
    "Models por meio de Digital Twins Conversacionais. 2026. Monografia "
    "(Bacharelado em Sistemas de Informação) – Escola de Artes, Ciências e "
    "Humanidades, Universidade de São Paulo, São Paulo, 2026."
)

RESUMO = (
    "Modelos de linguagem de grande porte (LLMs) produzem respostas variáveis "
    "para uma mesma entrada, dependendo de ajustes internos de geração e da "
    "instrução fornecida ao sistema. Avaliar essa variabilidade torna-se "
    "importante à medida que esses modelos passam a ser usados em aplicações "
    "sensíveis. Este trabalho propõe e valida uma prova de conceito de gêmeos "
    "digitais conversacionais: em um ponto fixo do diálogo, cópias paralelas de "
    "uma mesma conversa recebem pequenas alterações controladas — nos parâmetros "
    "de geração ou no nível de detalhamento da instrução — e têm seus desfechos "
    "comparados por métricas automáticas, por um modelo de linguagem usado como "
    "avaliador e por revisão humana. Um estudo piloto confirmou a viabilidade "
    "técnica e o custo baixo da abordagem, e um estudo principal, mais amplo e "
    "repetido cinco vezes, aprofundou a análise. O principal resultado é que "
    "boa parte da variação observada entre as respostas decorre do acaso "
    "natural do modelo (cerca de 20%), e não das alterações intencionais (cerca "
    "de 9%) — o que reforça a necessidade de uma réplica de controle na "
    "avaliação. A revisão humana mostrou ainda que o avaliador automático é "
    "confiável para medir a correção factual, mas tende a superestimar a "
    "estabilidade do conteúdo. Conclui-se que a abordagem é viável e "
    "reprodutível como ponto de partida metodológico para o monitoramento da "
    "variabilidade de assistentes conversacionais, com limitações relacionadas "
    "ao uso de um usuário simulado e ao tamanho da amostra."
)

PALAVRAS_CHAVE = (
    "Palavras-chave: Modelos de Linguagem de Grande Porte. Gêmeos Digitais. "
    "Avaliação de IA Generativa. Engenharia de Prompt. Experimentação Controlada."
)

ABSTRACT_REF = (
    "DIONÍSIO, Henrique Costa. Evaluation of Variability in Large Language "
    "Models through Conversational Digital Twins. 2026. Monograph (Bachelor of "
    "Information Systems) – School of Arts, Sciences and Humanities, University "
    "of São Paulo, São Paulo, 2026."
)

ABSTRACT = (
    "Large language models (LLMs) produce variable outputs for the same input, "
    "depending on internal generation settings and the instruction given to the "
    "system. Evaluating this variability becomes important as these models are "
    "adopted in sensitive applications. This work proposes and validates a proof "
    "of concept of conversational digital twins: at a fixed dialogue point, "
    "parallel copies of the same conversation receive small controlled changes — "
    "in generation parameters or in the level of detail of the instruction — and "
    "their outcomes are compared through automatic metrics, a language model used "
    "as a judge, and human review. A pilot study confirmed the technical "
    "feasibility and low cost of the approach, and a larger main study, repeated "
    "five times, deepened the analysis. The main finding is that much of the "
    "observed variation between responses stems from the model's natural "
    "randomness (about 20%) rather than from the intentional changes (about 9%), "
    "which reinforces the need for a control replicate in the evaluation. Human "
    "review further showed that the automatic judge is reliable for measuring "
    "factual correctness but tends to overestimate content stability. The "
    "approach is viable and reproducible as a methodological starting point for "
    "monitoring variability in conversational assistants, with limitations "
    "related to user simulation and sample size."
)

KEYWORDS = (
    "Keywords: Large Language Models. Digital Twins. Generative AI Evaluation. "
    "Prompt Engineering. Controlled Experimentation."
)

AGRADECIMENTOS = (
    "À minha orientadora, Profa. Dra. Sarajane Marques Peres, pela orientação "
    "atenta e pelas valiosas contribuições ao longo do trabalho. À Universidade "
    "de São Paulo e à Escola de Artes, Ciências e Humanidades, pela formação. "
    "Aos recursos de pesquisa (FAPESP/OpenAI) que viabilizaram a "
    "experimentação. À minha família e amigos, pelo apoio constante."
)

SIGLAS = [
    ("API", "Interface de Programação de Aplicações (Application Programming Interface)"),
    ("CLI", "Interface de Linha de Comando (Command Line Interface)"),
    ("EACH", "Escola de Artes, Ciências e Humanidades"),
    ("GPT", "Generative Pre-trained Transformer"),
    ("LLM", "Modelo de Linguagem de Grande Porte (Large Language Model)"),
    ("ORM", "Mapeamento Objeto-Relacional (Object-Relational Mapping)"),
    ("POC", "Prova de Conceito"),
    ("USP", "Universidade de São Paulo"),
]

FIGURAS = [
    "Pipeline de gêmeos digitais conversacionais",
    "Variação natural entre referência e réplica idêntica, por execução",
    "Efeito atribuível médio por tipo de perturbação",
    "Factualidade média por nível de instrução",
    "Robustez média por combinação de parâmetros",
    "Taxa de resolução por tipo de gêmeo",
    "Robustez: notas humanas vs. notas do juiz automático",
    "Indicadores médios por tipo de conversa",
]

TABELAS = [
    "Os nove gêmeos criados em cada ponto de cópia",
    "Comparação entre o piloto e o estudo principal",
    "Métricas utilizadas",
    "Variação natural e efeito atribuível (estudo principal)",
    "Avaliação humana vs. automática",
    "Indicadores médios por tipo de conversa (juiz automático)",
]

# Parágrafos pré-textuais a remover (por prefixo do texto).
DELETE_PREFIXES = [
    "(sugestão: use o recurso de inserir Índice",
    "(lista elaborada seguindo a mesma ordem",
    "(item opcional, inclua se necessário",
    "(listar em ordem alfabética)",
    "(nomes de instituições, mesmo se em língua estrangeira",
    "Lista de símbolos",
    "GHz",
    "Mpx",
    "(opcional, inclua se necessário por haver muitos símbolos)",
    "Lista de algoritmos, de quadros",
    "(listas opcionais, inclua se necessário",
    "(recomendamos que você adote o uso de Estilos",
    "Escreva aqui a sua dedicatória",
    "Texto de exemplo, texto de exemplo",
    "“Escreva aqui uma epígrafe",
    "“Escreva aqui uma epígrafe",
    "(Autor da epígrafe)",
    "Coorientador/a:",
    "ABNT \t",
    "BI \t",
    "ISO \t",
    "ULM \t",
]


def _resolve_image_path(raw: str):
    path = Path(raw.strip())
    if not path.is_absolute():
        path = (MD_PATH.parent / path).resolve()
        if not path.exists():
            path = (ROOT / raw.strip().lstrip("./")).resolve()
    return path if path.exists() else None


def _add_formatted_run(paragraph, text: str, bold: bool = False, italic: bool = False) -> None:
    run = paragraph.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")


def _add_markdown_paragraph(doc: Document, line: str, style: str = "Normal") -> None:
    p = doc.add_paragraph(style=style)
    pattern = re.compile(r"(\*\*[^*]+\*\*|\*[^*]+\*)")
    for part in pattern.split(line):
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            _add_formatted_run(p, part[2:-2], bold=True)
        elif part.startswith("*") and part.endswith("*"):
            _add_formatted_run(p, part[1:-1], italic=True)
        else:
            _add_formatted_run(p, part)


def _insert_figure(doc: Document, image_path, caption: str) -> None:
    global FIGURE_COUNTER
    FIGURE_COUNTER += 1
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(str(image_path), width=Cm(14))
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _add_formatted_run(cap, f"Figura {FIGURE_COUNTER} – {caption}")


def _parse_table_block(lines):
    rows = [line.strip() for line in lines if line.strip()]
    if len(rows) < 2:
        return [], []
    header = [c.strip() for c in rows[0].strip("|").split("|")]
    body = [[c.strip() for c in row.strip("|").split("|")] for row in rows[2:]]
    return header, body


def _add_table(doc: Document, header, body) -> None:
    if not header:
        return
    table = doc.add_table(rows=1 + len(body), cols=len(header))
    table.style = "Table Grid"
    for j, text in enumerate(header):
        table.rows[0].cells[j].text = text
    for i, row in enumerate(body, start=1):
        for j, text in enumerate(row):
            if j < len(table.rows[i].cells):
                table.rows[i].cells[j].text = text
    doc.add_paragraph()


def _process_codeblock(doc: Document, lines) -> None:
    p = doc.add_paragraph(style="Normal")
    run = p.add_run("\n".join(lines))
    run.font.name = "Courier New"
    run.font.size = Pt(9)


def append_markdown_body(doc: Document, md_text: str) -> None:
    """Anexa o corpo (a partir de '## 1. Introdução') ao documento, com figuras e tabelas."""
    global FIGURE_COUNTER
    FIGURE_COUNTER = 0
    lines = md_text.splitlines()
    i = 0
    in_code = False
    code_buf = []
    table_buf = []
    in_table = False
    started = False
    image_pattern = re.compile(r"!\[([^\]]*)\]\(([^)]+)\)")
    fig_caption_pattern = re.compile(r"^\*\*Figura\s+(\d+)\*\*\s*[–-]\s*(.+)$")

    while i < len(lines):
        line = lines[i]
        if not started:
            if line.startswith("## 1. Introdução"):
                started = True
            else:
                i += 1
                continue
        if line.strip().startswith("```"):
            if in_code:
                _process_codeblock(doc, code_buf)
                code_buf = []
                in_code = False
            else:
                in_code = True
            i += 1
            continue
        if in_code:
            code_buf.append(line)
            i += 1
            continue
        if line.strip().startswith("|"):
            in_table = True
            table_buf.append(line)
            i += 1
            continue
        elif in_table:
            _add_table(doc, *_parse_table_block(table_buf))
            table_buf = []
            in_table = False
            continue
        if line.strip() == "---":
            i += 1
            continue
        if line.startswith("## "):
            doc.add_heading(line[3:].strip(), level=1)
            i += 1
            continue
        if line.startswith("### "):
            doc.add_heading(line[4:].strip(), level=2)
            i += 1
            continue
        if line.startswith("#### "):
            doc.add_heading(line[5:].strip(), level=3)
            i += 1
            continue
        m_cap = fig_caption_pattern.match(line.strip())
        if m_cap:
            caption = m_cap.group(2).strip()
            i += 1
            while i < len(lines) and not lines[i].strip():
                i += 1
            if i < len(lines):
                m_img = image_pattern.search(lines[i])
                if m_img:
                    img_path = _resolve_image_path(m_img.group(2))
                    if img_path:
                        _insert_figure(doc, img_path, caption)
                    else:
                        _add_markdown_paragraph(doc, f"[Figura ausente: {m_img.group(2)}]")
                    i += 1
            continue
        m_img = image_pattern.search(line)
        if m_img and line.strip().startswith("!["):
            img_path = _resolve_image_path(m_img.group(2))
            if img_path:
                _insert_figure(doc, img_path, m_img.group(1) or "Figura")
            i += 1
            continue
        if line.strip():
            _add_markdown_paragraph(doc, line.strip())
        i += 1

    if in_table and table_buf:
        _add_table(doc, *_parse_table_block(table_buf))


def set_text(p: Paragraph, text: str) -> None:
    """Substitui o texto do parágrafo preservando a formatação do 1º run."""
    if p.runs:
        p.runs[0].text = text
        for r in p.runs[1:]:
            r._element.getparent().remove(r._element)
    else:
        run = p.add_run(text)
        run.font.name = "Times New Roman"


def is_heading1(p: Paragraph) -> bool:
    return p.style is not None and p.style.name == "Heading 1"


def insert_paragraph_after(p: Paragraph, text: str = "", bold: bool = False) -> Paragraph:
    new_p = OxmlElement("w:p")
    p._p.addnext(new_p)
    para = Paragraph(new_p, p._parent)
    run = para.add_run(text)
    run.bold = bold
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    return para


def replace_first(doc: Document, needle: str, value: str) -> bool:
    for p in doc.paragraphs:
        if is_heading1(p):
            break
        if p.text.strip() == needle:
            set_text(p, value)
            return True
    return False


def replace_startswith(doc: Document, prefix: str, value: str) -> bool:
    for p in doc.paragraphs:
        if is_heading1(p):
            break
        if p.text.strip().startswith(prefix):
            set_text(p, value)
            return True
    return False


def replace_all_text(doc: Document, needle: str, value: str) -> None:
    for p in doc.paragraphs:
        if is_heading1(p):
            break
        if p.text.strip() == needle:
            set_text(p, value)


def fill_list_after(doc: Document, heading: str, entries: list[str]) -> None:
    """Insere `entries` (já formatadas) logo após o parágrafo do título `heading`."""
    anchor = None
    for p in doc.paragraphs:
        if is_heading1(p):
            break
        if p.text.strip() == heading:
            anchor = p
            break
    if anchor is None:
        return
    for entry in entries:
        anchor = insert_paragraph_after(anchor, entry)


def delete_pretextual_junk(doc: Document) -> None:
    to_delete = []
    for p in doc.paragraphs:
        if is_heading1(p):
            break
        txt = p.text.strip()
        if not txt:
            continue
        if txt in ("...", "…"):
            to_delete.append(p)
            continue
        for prefix in DELETE_PREFIXES:
            if txt.startswith(prefix):
                to_delete.append(p)
                break
    for p in to_delete:
        p._element.getparent().remove(p._element)


def insert_toc_after(doc: Document, heading: str) -> None:
    anchor = None
    for p in doc.paragraphs:
        if is_heading1(p):
            break
        if p.text.strip() == heading:
            anchor = p
            break
    if anchor is None:
        return
    new_p = OxmlElement("w:p")
    anchor._p.addnext(new_p)

    def _run(child):
        r = OxmlElement("w:r")
        r.append(child)
        new_p.append(r)

    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    _run(begin)
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = ' TOC \\o "1-3" \\h \\z \\u '
    _run(instr)
    sep = OxmlElement("w:fldChar")
    sep.set(qn("w:fldCharType"), "separate")
    _run(sep)
    t = OxmlElement("w:t")
    t.text = "Atualize o sumário no Word: clique com o botão direito e escolha “Atualizar campo”."
    _run(t)
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    _run(end)


def trim_template_body(doc: Document) -> None:
    """Remove o corpo de exemplo do template (do 1º Heading 1 até o fim,
    preservando o sectPr final)."""
    body = doc.element.body
    first_h1 = None
    for p in doc.paragraphs:
        if is_heading1(p):
            first_h1 = p._element
            break
    if first_h1 is None:
        return
    started = False
    for child in list(body):
        if child is first_h1:
            started = True
        if started and not child.tag.endswith("sectPr"):
            body.remove(child)


def main() -> None:
    doc = Document(str(TEMPLATE_PATH))

    # Capa e folha de rosto
    replace_all_text(doc, "AUTOR DO TRABALHO", AUTOR)
    replace_all_text(doc, "Título do trabalho: subtítulo do trabalho", TITULO)
    replace_all_text(doc, "2025", "2026")
    replace_startswith(doc, "Monografia apresentada",
                       "Monografia apresentada à Escola de Artes, Ciências e "
                       "Humanidades da Universidade de São Paulo, como parte dos "
                       "requisitos para obtenção do título de Bacharel em Sistemas "
                       "de Informação, no âmbito da disciplina ACH2018 – Projeto "
                       "Supervisionado ou de Graduação II.")
    replace_startswith(doc, "Modalidade:", "Modalidade: TCC Individual Semestral")
    replace_startswith(doc, "Orientador/a:",
                       "Orientadora: Profa. Dra. Sarajane Marques Peres")

    # Agradecimentos: usa o 1º parágrafo de exemplo como texto real
    replace_startswith(doc, "Texto de exemplo, texto de exemplo", AGRADECIMENTOS)
    replace_first(doc, "…", "")

    # Resumo / Abstract
    replace_startswith(doc, "SOBRENOME, Nome", RESUMO_REF)
    replace_startswith(doc, "Escreva aqui o texto do seu resumo", RESUMO)
    replace_startswith(doc, "Palavras-chaves:", PALAVRAS_CHAVE)
    replace_startswith(doc, "SURNAME, FirstName", ABSTRACT_REF)
    replace_startswith(doc, "Write here the English version", ABSTRACT)
    replace_startswith(doc, "Keywords:", KEYWORDS)

    # Listas pré-textuais
    fill_list_after(doc, "Lista de figuras",
                    [f"Figura {i} – {c}" for i, c in enumerate(FIGURAS, 1)])
    fill_list_after(doc, "Lista de tabelas",
                    [f"Tabela {i} – {c}" for i, c in enumerate(TABELAS, 1)])
    fill_list_after(doc, "Lista de abreviaturas e siglas",
                    [f"{s}\t\t{d}" for s, d in SIGLAS])
    insert_toc_after(doc, "Sumário")

    # Limpa textos de exemplo restantes (dedicatória, epígrafe, notas, etc.)
    delete_pretextual_junk(doc)

    # Remove o corpo de exemplo e anexa o conteúdo real
    trim_template_body(doc)
    md_text = MD_PATH.read_text(encoding="utf-8")
    append_markdown_body(doc, md_text)

    OUTPUT_PATH.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(OUTPUT_PATH))
    print(f"Documento salvo: {OUTPUT_PATH}")
    print(f"Figuras embutidas: {FIGURE_COUNTER}")


if __name__ == "__main__":
    main()
